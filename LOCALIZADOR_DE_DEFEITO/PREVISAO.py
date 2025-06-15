import numpy as np
import matplotlib.pyplot as plt
import os
import csv
import re
import pandas as pd
from sklearn.ensemble import RandomForestRegressor
from sklearn.model_selection import train_test_split
from sklearn.metrics import mean_absolute_error, r2_score
from colorama import init, Fore
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

init(autoreset=True)


# -----------------------------
# Funções para gerar dados MCD
# -----------------------------
def ler_arquivo_com_lista(caminho_arquivo):
    with open(caminho_arquivo, 'r') as f:
        linhas = f.readlines()
    dados = [float(linha.strip()) for linha in linhas[1:] if linha.strip()]
    return dados


def calcular_curvatura_modal(UY, h):
    curvatura = np.zeros_like(UY)
    for i in range(1, len(UY) - 1):
        curvatura[i] = (UY[i + 1] - 2 * UY[i] + UY[i - 1]) / (h ** 2)
    curvatura[0] = curvatura[1]
    curvatura[-1] = curvatura[-2]
    return curvatura


def calcular_MCD(curvatura_1, curvatura_2):
    return np.abs(np.abs(curvatura_1) - np.abs(curvatura_2))


def gerar_dados_mcd(pasta_arquivos, pasta_resultados):
    os.makedirs(pasta_resultados, exist_ok=True)

    s = np.linspace(0, 100, 101)
    h = 100 / 100

    caminho_sem_defeito = os.path.join(pasta_arquivos, 'sem_defeito.txt')
    UY_sem_defeito = np.array(ler_arquivo_com_lista(caminho_sem_defeito))
    curvatura_sem_defeito = calcular_curvatura_modal(UY_sem_defeito, h)

    todas_colunas_mcd = []
    cabecalho = []

    padrao = r"VETOR_MODAL_(\d+)_(\d+\.\d+)_(\d+\.\d+)_(\d+\.\d+)\.txt"

    for nome_arquivo in os.listdir(pasta_arquivos):
        if nome_arquivo.startswith("VETOR_MODAL") and nome_arquivo.endswith(".txt"):
            match = re.match(padrao, nome_arquivo)
            if match:
                valor1, valor2, valor3, valor4 = match.groups()
                caminho_arquivo = os.path.join(pasta_arquivos, nome_arquivo)
                try:
                    UY_defeito = np.array(ler_arquivo_com_lista(caminho_arquivo))
                    curvatura_defeito = calcular_curvatura_modal(UY_defeito, h)
                    mcd = calcular_MCD(curvatura_sem_defeito, curvatura_defeito)

                    todas_colunas_mcd.append(mcd)
                    cabecalho.append(f'MCD_{valor1}_{valor2}_{valor3}_{valor4}')

                    plt.figure(figsize=(8, 4))
                    plt.plot(s, mcd)
                    plt.xlabel('Posição (%)')
                    plt.ylabel('MCD')
                    plt.title(f'MCD - {nome_arquivo}')
                    plt.grid(True)
                    plt.tight_layout()
                    plt.savefig(os.path.join(pasta_resultados, f"MCD_{valor1}_{valor2}_{valor3}_{valor4}.png"))
                    plt.close()

                except FileNotFoundError:
                    print(f"Arquivo {nome_arquivo} não encontrado.")

    caminho_csv_unico = os.path.join(pasta_resultados, "MCD_BANCO_DE_DADOS.csv")
    with open(caminho_csv_unico, 'w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(cabecalho)
        for i in range(len(s)):
            linha = [col[i] for col in todas_colunas_mcd]
            writer.writerow(linha)
    print(Fore.GREEN + f"Arquivo CSV com MCDs salvo em: {caminho_csv_unico}")
    return caminho_csv_unico


# -----------------------------
# Funções para ML e relatório
# -----------------------------
def extrair_alvos(nome_coluna):
    partes = nome_coluna.split("_")
    posicao = float(partes[2])
    profundidade = float(partes[3])
    espessura = float(partes[4])
    return posicao, profundidade, espessura


def avaliar_modelo(y_true, y_pred, y_train_true, y_train_pred, nome, doc):
    erro_teste = mean_absolute_error(y_true, y_pred)
    erro_treino = mean_absolute_error(y_train_true, y_train_pred)
    r2_teste = r2_score(y_true, y_pred)
    r2_treino = r2_score(y_train_true, y_train_pred)

    print(Fore.CYAN + f"\n--- {nome.upper()} ---")
    print(Fore.YELLOW + f"Erro médio (TREINO): {erro_treino:.4f}")
    print(Fore.GREEN + f"Erro médio (TESTE) : {erro_teste:.4f}")
    print(Fore.YELLOW + f"R² (TREINO)         : {r2_treino:.4f}")
    print(Fore.GREEN + f"R² (TESTE)          : {r2_teste:.4f}")
    print(Fore.CYAN + "-" * 40)

    doc.add_heading(f'Resultados para {nome}', level=2)
    doc.add_paragraph(f"Erro médio (TREINO): {erro_treino:.4f}")
    doc.add_paragraph(f"Erro médio (TESTE) : {erro_teste:.4f}")
    doc.add_paragraph(f"R² (TREINO)        : {r2_treino:.4f}")
    doc.add_paragraph(f"R² (TESTE)         : {r2_teste:.4f}")

    return erro_teste, r2_teste


def prever_defeito(coluna_dados, modelos):
    X_novo = np.array(coluna_dados).reshape(1, -1)
    return {
        "posicao": modelos["posicao"].predict(X_novo)[0],
        "profundidade": modelos["profundidade"].predict(X_novo)[0],
        "espessura": modelos["espessura"].predict(X_novo)[0]
    }


# -----------------------------
# Fluxo principal
# -----------------------------
def main():
    pasta_arquivos = r'C:\Users\GVA\Desktop\dico'  # ajuste conforme sua pasta
    pasta_resultados = 'resultados_MCD'

    # 1) Gerar CSV com dados MCD
    caminho_csv = gerar_dados_mcd(pasta_arquivos, pasta_resultados)

    # 2) Carregar dados para ML
    dados = pd.read_csv(caminho_csv)
    X = dados.T
    y = pd.DataFrame([extrair_alvos(col) for col in dados.columns], columns=["posicao", "profundidade", "espessura"])

    # 3) Dividir treino/teste
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    # 4) Treinar modelos RF
    modelos = {}
    importancias = {}

    doc = Document()
    titulo = doc.add_heading('Relatório de Previsão de Defeitos - Random Forest', level=1)
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('Introdução', level=1)
    doc.add_paragraph(
        "Este relatório apresenta os resultados obtidos com o uso do algoritmo Random Forest para prever a localização, profundidade e espessura de defeitos em materiais com base em dados extraídos de sensores.")

    doc.add_heading('Treinamento dos Modelos', level=1)

    for alvo in y.columns:
        modelo = RandomForestRegressor(n_estimators=100, random_state=42)
        modelo.fit(X_train, y_train[alvo])
        modelos[alvo] = modelo

        y_pred = modelo.predict(X_test)
        y_train_pred = modelo.predict(X_train)

        erro, r2 = avaliar_modelo(y_test[alvo], y_pred, y_train[alvo], y_train_pred, alvo, doc)

        importancias[alvo] = modelo.feature_importances_

        fig, ax = plt.subplots()
        ax.bar(range(len(modelo.feature_importances_)), modelo.feature_importances_)
        ax.set_title(f'Importância das features - {alvo}')
        ax.set_xlabel("Entradas")
        ax.set_ylabel("Importância")
        fig.tight_layout()
        fig_path = os.path.join(pasta_resultados, f"importancia_{alvo}.png")
        fig.savefig(fig_path)
        plt.close(fig)
        doc.add_picture(fig_path, width=Inches(5.5))

    media_importancia = np.mean([importancias[a] for a in importancias], axis=0)
    importancia_ordenada = sorted(zip(media_importancia, X.columns), reverse=True)[:10]
    doc.add_heading("Importância Global das Features", level=1)
    for i, (valor, idx) in enumerate(importancia_ordenada, 1):
        doc.add_paragraph(f"{i}. Entrada {idx} -> Importância Média: {valor:.4f}")

    doc.add_heading("Comparacão Real vs Predito (Teste)", level=1)
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Exemplo'
    hdr_cells[1].text = 'Alvo'
    hdr_cells[2].text = 'Real'
    hdr_cells[3].text = 'Predito'

    for i in range(min(10, len(y_test))):
        for alvo in y.columns:
            row_cells = tabela.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = alvo
            row_cells[2].text = f"{y_test.iloc[i][alvo]:.4f}"
            row_cells[3].text = f"{modelos[alvo].predict([X_test.iloc[i]])[0]:.4f}"

    for alvo in y.columns:
        fig, ax = plt.subplots()
        ax.scatter(y_test[alvo], modelos[alvo].predict(X_test), color='blue', alpha=0.6)
        ax.plot([y_test[alvo].min(), y_test[alvo].max()], [y_test[alvo].min(), y_test[alvo].max()], 'r--')
        ax.set_xlabel("Valor real")
        ax.set_ylabel("Valor predito")
        ax.set_title(f'Real vs Predito - {alvo}')
        fig.tight_layout()
        fig_path = os.path.join(pasta_resultados, f"real_vs_predito_{alvo}.png")
        fig.savefig(fig_path)
        plt.close(fig)
        doc.add_picture(fig_path, width=Inches(5.5))

    # Previsão de novo vetor (exemplo)
    coluna_fixa = [
        1.09E-06, 1.09E-06, 1.77E-07, 3.10E-08, 7.30E-08, 2.16E-07, 3.29E-07, 4.39E-07, 8.00E-06,
        2.23E-05, 0.000141506, 2.88E-05, 8.19E-06, 6.39E-07, 3.62E-07, 3.72E-07, 3.55E-07, 3.52E-07,
        3.59E-07, 3.73E-07, 3.25E-06, 1.24E-07, 1.55E-07, 1.89E-07, 2.22E-07, 2.58E-07, 2.94E-07,
        3.31E-07, 3.64E-07, 4.04E-07, 4.40E-07, 4.77E-07, 5.15E-07, 5.53E-07, 5.90E-07, 6.29E-07,
        6.64E-07, 7.02E-07, 7.40E-07, 7.74E-07, 6.51E-06, 9.41E-07, 9.61E-07, 9.76E-07, 9.91E-07,
        1.01E-06, 1.02E-06, 1.03E-06, 1.04E-06, 1.05E-06, 1.06E-06, 1.06E-06, 1.07E-06, 1.07E-06,
        1.08E-06, 1.08E-06, 1.08E-06, 1.07E-06, 1.07E-06, 1.07E-06, 4.52E-06, 1.18E-06, 1.19E-06,
        1.19E-06, 1.19E-06, 1.20E-06, 1.19E-06, 1.19E-06, 1.18E-06, 1.18E-06, 1.17E-06, 1.16E-06,
        1.15E-06, 1.13E-06, 1.11E-06, 1.09E-06, 1.07E-06, 1.05E-06, 1.01E-06, 7.56E-07, 9.17E-06,
        6.37E-06, 6.30E-06, 2.41E-06, 6.24E-07, 5.95E-07, 5.68E-07, 5.43E-07, 5.16E-07, 4.89E-07,
        4.65E-07, 4.39E-07, 4.14E-07, 3.90E-07, 3.62E-07, 3.37E-07, 3.12E-07, 2.69E-07, 2.52E-07,
        5.70E-07, 5.70E-07]

    if len(coluna_fixa) != X.shape[1]:
        raise ValueError(f"A coluna fornecida tem {len(coluna_fixa)} valores, mas o modelo espera {X.shape[1]}.")

    resultado = prever_defeito(coluna_fixa, modelos)
    doc.add_heading("Previsão para Novo Vetor", level=1)
    doc.add_paragraph(f"Posição prevista     : {resultado['posicao']:.2f} %")
    doc.add_paragraph(f"Profundidade prevista: {resultado['profundidade']:.2f} %")
    doc.add_paragraph(f"Espessura prevista   : {resultado['espessura']:.4f} m")

    doc.add_heading("Conclusão", level=1)
    doc.add_paragraph(
        "O modelo apresentou bom desempenho geral, especialmente na predição da profundidade. A espessura, no entanto, não demonstrou correlação significativa, indicando necessidade de dados mais robustos para esse alvo.")

    arquivo_relatorio = os.path.join(pasta_resultados, "Relatorio_Previsao_Defeitos.docx")
    doc.save(arquivo_relatorio)
    print(Fore.LIGHTMAGENTA_EX + f"\nRelatório gerado com sucesso: {arquivo_relatorio}")


if __name__ == "__main__":
    main()
